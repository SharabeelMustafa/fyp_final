from docx import Document

# Create a new Word document
doc = Document()

# Add title to the document
doc.add_heading("Gantt Chart for Campus Wheel Comsats", level=1)

# Define data for the Gantt chart
headers = ["Week", "Tasks by Muhammad Qasim Khan", "Tasks by Muhammad Shahbaz", "Tasks by Sharhabeel Mustafa Mehmood"]
data = [
    ["1", "System Design", "System Design", "Backend Development"],
    ["2", "System Design", "System Design", "Backend Development"],
    ["3", "System Design", "App Frontend Development", "Backend Development"],
    ["4", "Web Front End Development", "App Frontend Development", "Database Setup"],
    ["5", "Web Front End Development", "Backend Development, Database", "API Development"],
    ["6", "Web Front End Development", "Backend Development", "API Development"],
    ["7", "Testing and Validation", "Frontend Enhancements", "Data Management"],
    ["8", "Testing and Validation", "UI/UX Enhancements", "Final Database Testing"],
    ["9", "GPS Tracking Integration", "Payment Integration", "Update Backend for Payment"],
    ["10", "Real-Time Tracking Testing", "Secure Payment Testing", "GPS Tracking Testing"],
    ["11", "Final System Review", "Final System Review", "Final System Review"],
    ["12", "Project Refinement", "Project Refinement", "Project Refinement"],
]

# Add table to the document
table = doc.add_table(rows=len(data) + 1, cols=len(headers))
table.style = 'Table Grid'

# Add headers to the table
for col_idx, header in enumerate(headers):
    table.cell(0, col_idx).text = header

# Add data to the table
for row_idx, row_data in enumerate(data, start=1):
    for col_idx, cell_data in enumerate(row_data):
        table.cell(row_idx, col_idx).text = cell_data

# Save the document
file_path = "Campus_Wheel_Gantt_Chart.docx"
doc.save(file_path)

print(f"Document saved as {file_path}")
